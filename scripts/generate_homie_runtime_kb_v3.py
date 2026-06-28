"""
Generate Homie Runtime KB v3 workbook and Developer Brief docx.

Usage:
  pip install openpyxl python-docx
  python scripts/generate_homie_runtime_kb_v3.py

Output:
  output/Homie_KB_Runtime_Phase1_4Menu_ManualAligned_v3.xlsx
  output/Homie_Developer_Brief_Phase1_4Menu_v3.docx

Governance:
  - Runtime KB must be clean.
  - AI must not invent lending criteria.
  - Customer data must be separated from knowledge base.
"""

from __future__ import annotations

from pathlib import Path
from datetime import date
from typing import Dict, List, Any

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.utils import get_column_letter

try:
    from docx import Document
    from docx.shared import Pt
except Exception:  # pragma: no cover
    Document = None

OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)

VERSION = "v3.0"
PHASE = "Phase1"
REVIEWED_BY = "Auriporn Chatthai / Homie Project Team"
REVIEW_DATE = date.today().isoformat()
COMMON = ["menu_scope", "is_active", "phase", "review_status", "source_reference", "version", "last_reviewed_by", "last_reviewed_at"]


def common(menu_scope: str, review_status: str = "Approved", source: str = "Homie Data Blueprint v3") -> List[Any]:
    return [menu_scope, True, PHASE, review_status, source, VERSION, REVIEWED_BY, REVIEW_DATE]


SHEETS: Dict[str, Dict[str, Any]] = {}


def add_sheet(name: str, headers: List[str], rows: List[List[Any]]) -> None:
    SHEETS[name] = {"headers": headers, "rows": rows}


# 00 Config
add_sheet(
    "00_Config",
    ["config_key", "config_value", "description"] + COMMON,
    [
        ["runtime_kb_name", "Homie_KB_Runtime_Phase1_4Menu_ManualAligned_v3", "Runtime KB name"] + common("CORE"),
        ["phase", PHASE, "Phase used by runtime reader"] + common("CORE"),
        ["read_rule", "is_active=TRUE AND phase=Phase1 AND review_status=Approved", "Mandatory runtime filter"] + common("CORE"),
        ["customer_pdf_internal_data_allowed", False, "Customer PDF must exclude internal data"] + common("PDF"),
        ["ai_can_invent_lending_rules", False, "AI must not invent lending criteria"] + common("GUARDRAIL"),
        ["customer_data_in_runtime_kb_allowed", False, "Customer data must be kept outside Runtime KB"] + common("GUARDRAIL"),
    ],
)

# 01 Menu Master
add_sheet(
    "01_Menu_Master",
    ["menu_id", "menu_name", "menu_description", "display_order", "icon_name"] + COMMON,
    [
        ["MENU01", "คำนวณสินเชื่ออาชีพอิสระ", "ประเมินรายได้สุทธิ ภาระหนี้ และเงินงวดเบื้องต้น", 1, "calculator"] + common("CALC"),
        ["MENU02", "Checklist", "สร้างรายการเอกสารตามกลุ่มผู้กู้ อาชีพ และวัตถุประสงค์การกู้", 2, "checklist"] + common("CHECKLIST"),
        ["MENU03", "Self Learning", "คู่มือเร็ว ข้อผิดพลาดที่พบบ่อย Infographic และ Quiz", 3, "learning"] + common("LEARNING"),
        ["MENU04", "Q&A", "ค้นหาคำตอบจากฐานความรู้ที่ผ่านการตรวจสอบ", 4, "qa"] + common("QA"),
    ],
)

# Guardrails
guardrails = [
    ["G001", "SYSTEM", "Runtime KB ต้องสะอาด", "อ่านเฉพาะข้อมูล Approved เท่านั้น", "critical", "filter_approved_only"],
    ["G002", "SYSTEM", "AI ห้ามเดาข้อมูลสินเชื่อเอง", "ถ้าไม่มีข้อมูล Approved ให้ fallback/escalate", "critical", "retrieve_only"],
    ["G003", "SYSTEM", "ข้อมูลลูกค้าต้องแยกจากฐานความรู้", "ห้ามเก็บข้อมูลลูกค้าใน Runtime KB", "critical", "separate_operational_data"],
    ["G004", "CALC", "ไม่ใช่ผลอนุมัติสินเชื่อ", "ทุกผลคำนวณต้องแสดง Disclaimer", "high", "append_disclaimer"],
    ["G005", "CHECKLIST", "Customer PDF ห้ามแสดงข้อมูลภายใน", "ซ่อน Risk Flag, Interview Focus, Staff Note, Source Reference", "critical", "hide_internal_fields"],
    ["G006", "QA", "ถ้าไม่พบคำตอบในฐาน Approved", "ใช้ fallback ห้ามแต่งคำตอบเอง", "critical", "fallback"],
]
add_sheet(
    "02_Response_Guardrails",
    ["guardrail_id", "applies_to", "rule_name", "rule_text", "severity", "action"] + COMMON,
    [row + common("GUARDRAIL") for row in guardrails],
)

# Runtime Lists
runtime_lists = []
for i, value in enumerate(["CORE", "CALC", "CHECKLIST", "LEARNING", "QA", "GUARDRAIL", "PDF", "BRANCH", "VALIDATION"], 1):
    runtime_lists.append(["menu_scope", value, value, i])
for i, value in enumerate(["Draft", "Need Review", "Approved", "Inactive"], 1):
    runtime_lists.append(["review_status", value, value, i])
for i, value in enumerate(["Phase1", "Phase2"], 1):
    runtime_lists.append(["phase", value, value, i])
for i, value in enumerate(["พนักงานสินเชื่อ", "หัวหน้างานสินเชื่อ", "ผู้ช่วยหัวหน้าส่วน", "Supervisor", "Admin", "Tester / Pilot Phase 1"], 1):
    runtime_lists.append(["staff_role", value, value, i])
for i, value in enumerate(["รอติดต่อ", "แจ้งเอกสารแล้ว", "รอเอกสาร", "รับเอกสารแล้ว", "ปิดเคส", "ส่งต่อผู้รับผิดชอบ"], 1):
    runtime_lists.append(["follow_up_status", value, value, i])
for i, value in enumerate(["Required", "Conditional", "Optional"], 1):
    runtime_lists.append(["required_level", value, value, i])
add_sheet("03_Runtime_Lists", ["list_name", "list_value", "display_text", "sort_order"] + COMMON, [row + common("CORE") for row in runtime_lists])

# Branch Master
branches = [
    ["BR001", "สาขานครราชสีมา", "044 248 201 (02-07)", "เขตนครราชสีมา", "ภาคตะวันออกเฉียงเหนือตอนล่าง"],
    ["BR002", "สาขาถนนจอมพล", "044 269 205 (06-07)", "เขตนครราชสีมา", "ภาคตะวันออกเฉียงเหนือตอนล่าง"],
    ["BR003", "สาขาบัวใหญ่", "044 913 921 (22), 044 913 933 (34)", "เขตนครราชสีมา", "ภาคตะวันออกเฉียงเหนือตอนล่าง"],
    ["BR004", "สาขาเซ็นทรัล โคราช", "044 229 457 (58-62)", "เขตนครราชสีมา", "ภาคตะวันออกเฉียงเหนือตอนล่าง"],
    ["BR005", "สาขาปากช่อง", "044 316 395 (96-98)", "เขตนครราชสีมา", "ภาคตะวันออกเฉียงเหนือตอนล่าง"],
    ["BR006", "สาขาชัยภูมิ", "044 813 605 (07,09-12)", "เขตนครราชสีมา", "ภาคตะวันออกเฉียงเหนือตอนล่าง"],
    ["BR007", "สาขาภูเขียว", "044 861 756 (57), 044 861 845 (46)", "เขตนครราชสีมา", "ภาคตะวันออกเฉียงเหนือตอนล่าง"],
    ["BR008", "สาขาบุรีรัมย์", "044 616 972 (73-75)", "เขตนครราชสีมา", "ภาคตะวันออกเฉียงเหนือตอนล่าง"],
    ["BR009", "สาขานางรอง", "044 624 205 (06-08)", "เขตนครราชสีมา", "ภาคตะวันออกเฉียงเหนือตอนล่าง"],
    ["BR010", "สาขาอุบลราชธานี", "045 265 801 (02-08)", "เขตอุบลราชธานี", "ภาคตะวันออกเฉียงเหนือตอนล่าง"],
    ["BR011", "สาขาเซ็นทรัล อุบล", "045 422 381 (82-84)", "เขตอุบลราชธานี", "ภาคตะวันออกเฉียงเหนือตอนล่าง"],
    ["BR012", "สาขาศรีสะเกษ", "045 617 742 (43-45,47-48)", "เขตอุบลราชธานี", "ภาคตะวันออกเฉียงเหนือตอนล่าง"],
    ["BR013", "สาขาอำนาจเจริญ", "045 511 367 (68-70)", "เขตอุบลราชธานี", "ภาคตะวันออกเฉียงเหนือตอนล่าง"],
    ["BR014", "สาขาสุรินทร์", "044 511 967 (68-70)", "เขตอุบลราชธานี", "ภาคตะวันออกเฉียงเหนือตอนล่าง"],
    ["BR015", "สาขายโสธร", "045 711 740, 045 714 592-3", "เขตอุบลราชธานี", "ภาคตะวันออกเฉียงเหนือตอนล่าง"],
]
add_sheet("04_Branch_Master", ["branch_id", "branch_name", "branch_phone", "branch_group", "region"] + COMMON, [row + common("BRANCH", source="User-provided branch phone list") for row in branches])

# Calculation input schema
calc_inputs = [
    ["IN001", "customer_name", "ชื่อลูกค้า / ผู้มาติดต่อ", "text", True, ""],
    ["IN002", "contact_info", "เบอร์ติดต่อ / LINE ID", "text", False, "ห้ามนำไปเก็บใน Runtime KB"],
    ["IN003", "applicant_group_id", "กลุ่มผู้กู้", "dropdown", True, "AG01/AG02/AG03"],
    ["IN004", "occupation_id", "อาชีพอิสระ", "dropdown", True, "OCC001-OCC028"],
    ["IN005", "income_evidence_type", "หลักฐานรายได้", "dropdown", True, "Statement/dStatement/บัญชีรายรับรายจ่าย/สัญญา/อื่น ๆ"],
    ["IN006", "gross_monthly_income", "รายได้รวมเฉลี่ยต่อเดือน", "number", True, "บาท"],
    ["IN007", "expense_method", "วิธีหักค่าใช้จ่าย", "dropdown", True, "actual_or_standard"],
    ["IN008", "actual_business_expense", "ค่าใช้จ่ายกิจการจริงต่อเดือน", "number", False, "บาท"],
    ["IN009", "standard_expense_rate", "อัตราค่าใช้จ่ายมาตรฐานตามอาชีพ", "percent", "AUTO", "ดึงจาก 12_Expense_Assumptions_Occ28 เฉพาะ Approved"],
    ["IN010", "net_income_estimated", "รายได้สุทธิเบื้องต้น", "number", "AUTO", "คำนวณ"],
    ["IN011", "credit_bureau_debt", "ภาระหนี้จากเครดิตบูโร", "number", True, "บาท"],
    ["IN012", "other_debt", "ภาระหนี้อื่นที่ต้องนำมาหัก", "number", False, "บาท"],
    ["IN013", "requested_loan_amount", "วงเงินที่ต้องการประเมิน", "number", False, "บาท"],
    ["IN014", "loan_term_years", "ระยะเวลากู้", "number", True, "ปี"],
    ["IN015", "interest_rate_assumption", "อัตราดอกเบี้ยสมมติฐาน", "percent", True, "ต่อปี"],
    ["IN016", "estimated_installment", "เงินงวดประมาณการ", "number", "AUTO", "PMT หรือ staff-entered during demo"],
    ["IN017", "income_expense_note", "หมายเหตุรายได้ / ค่าใช้จ่าย", "textarea", False, "ไม่เก็บใน Runtime KB"],
]
add_sheet("10_Calc_Input_Schema", ["input_id", "field_name", "display_label", "input_type", "required", "help_text"] + COMMON, [row + common("CALC") for row in calc_inputs])

add_sheet(
    "11_Income_Rules_Basic",
    ["income_rule_id", "rule_name", "calculation_method", "output_note"] + COMMON,
    [
        ["IR001", "ใช้ค่าใช้จ่ายจริง", "net_income = gross_monthly_income - actual_business_expense", "ใช้เมื่อมีหลักฐานค่าใช้จ่ายชัดเจน"] + common("CALC"),
        ["IR002", "ใช้อัตราค่าใช้จ่ายมาตรฐาน", "expense = gross_monthly_income * default_expense_rate; net_income = gross - expense", "ใช้เฉพาะอัตราที่ Approved ใน Runtime KB"] + common("CALC"),
        ["IR003", "ไม่มีอัตราค่าใช้จ่าย Approved", "ask_staff_to_enter_actual_expense", "ห้าม AI เดาเปอร์เซ็นต์"] + common("CALC"),
    ],
)

# OCC28 placeholders — deliberately Need Review until mapped from official manual/source workbook
occ_rows = []
for i in range(1, 29):
    occ_id = f"OCC{i:03d}"
    occ_rows.append([occ_id, "รอ map จากคู่มือ/ไฟล์เดิม", "Independent Occupation", "", "", "", "", "Need Review"] + common("CHECKLIST", review_status="Need Review", source="Pending mapping from AI_Knowledge_Assistant workbook / official manual"))
add_sheet("21_Occupation_Master_28", ["occupation_id", "occupation_name", "occupation_group", "aliases", "business_characteristics", "income_documents", "risk_ref", "checklist_status"] + COMMON, occ_rows)

# Expense assumptions placeholders — no invented percentages
expense_rows = []
for i in range(1, 29):
    occ_id = f"OCC{i:03d}"
    expense_rows.append([f"EXP{i:03d}", "", occ_id, "รอ map จากคู่มือ/ไฟล์เดิม", "รอระบุ", None, None, None, "กรอกค่าใช้จ่ายจริงจนกว่าจะมีอัตรา Approved", "ห้าม AI เดา % ค่าใช้จ่าย"] + common("CALC", review_status="Need Review", source="Pending approved occupation expense rate"))
add_sheet("12_Expense_Assumptions_Occ28", ["assumption_id", "applicant_group_id", "occupation_id", "occupation_name", "cost_type", "expense_rate_min", "expense_rate_max", "default_expense_rate", "evidence_required", "staff_note"] + COMMON, expense_rows)

# Debt rules
add_sheet(
    "13_Debt_Obligation_Rules",
    ["debt_rule_id", "debt_item", "treatment", "explanation", "example"] + COMMON,
    [
        ["DR001", "ภาระหนี้จากเครดิตบูโร", "deduct", "ต้องนำมาหักในการประเมินเบื้องต้น", "ยอดผ่อนต่อเดือน"] + common("CALC"),
        ["DR002", "หนี้บัตรเครดิต", "deduct_if_in_credit_bureau", "หักตามข้อมูลเครดิตบูโร", "ยอดชำระบัตรเครดิต"] + common("CALC"),
        ["DR003", "ค่าไฟฟ้า", "not_debt", "ไม่คิดเป็นภาระหนี้", "รายการตัดบัญชีค่าไฟ"] + common("CALC"),
        ["DR004", "ค่าเบี้ยประกันชีวิต", "not_debt", "ไม่คิดเป็นภาระหนี้", "รายการประกันชีวิต"] + common("CALC"),
    ],
)

add_sheet(
    "14_Installment_Estimation_Rules",
    ["rule_id", "rule_name", "formula", "note"] + COMMON,
    [
        ["PMT001", "เงินงวดประมาณการ", "PMT(monthly_rate, term_months, requested_loan_amount)", "ใช้เพื่อประเมินเบื้องต้นเท่านั้น"] + common("CALC"),
        ["DSR001", "Estimated DSR", "(existing_debt + estimated_installment) / net_income_estimated", "ห้ามใช้เป็นผลอนุมัติ"] + common("CALC"),
    ],
)

add_sheet(
    "15_Calc_Output_Template",
    ["output_section_id", "section_title", "field_list", "customer_visible"] + COMMON,
    [
        ["CO001", "สรุปข้อมูลที่ใช้ประเมิน", "applicant_group, occupation, evidence", False] + common("CALC"),
        ["CO002", "รายได้สุทธิเบื้องต้น", "gross, expense, net_income", False] + common("CALC"),
        ["CO003", "ภาระหนี้", "credit_bureau_debt, other_debt", False] + common("CALC"),
        ["CO004", "ผลประเมินเบื้องต้น", "remaining_income, estimated_installment, estimated_dsr, risk_alert", False] + common("CALC"),
        ["CO005", "Disclaimer", "mandatory_disclaimer", True] + common("CALC"),
    ],
)

add_sheet(
    "16_Calc_Risk_Alert",
    ["risk_id", "condition", "alert_text", "severity"] + COMMON,
    [
        ["R001", "gross_income <= 0", "กรุณาระบุรายได้รวมเฉลี่ยต่อเดือน", "high"] + common("CALC"),
        ["R002", "expense_rate_missing", "ยังไม่มีอัตราค่าใช้จ่ายมาตรฐานที่ผ่านการตรวจสอบ กรุณากรอกค่าใช้จ่ายจริงตามหลักฐาน", "critical"] + common("CALC"),
        ["R003", "estimated_dsr_high", "Estimated DSR สูง ควรตรวจสอบรายได้และภาระหนี้เพิ่มเติมก่อนส่งเคส", "high"] + common("CALC"),
        ["R004", "weak_evidence", "หลักฐานรายได้ยังไม่ชัด ควรขอเอกสารเพิ่มเติม", "medium"] + common("CALC"),
    ],
)

# Applicant groups
add_sheet(
    "20_Applicant_Group_3Part",
    ["applicant_group_id", "applicant_group_name", "description", "decision_hint", "risk_note"] + COMMON,
    [
        ["AG01", "อาชีพอิสระทั่วไป / ไม่จดทะเบียน", "บุคคลธรรมดาประกอบอาชีพอิสระทั่วไป", "ใช้เมื่อไม่มีการจดทะเบียนกิจการ/นิติบุคคล", "ต้องตรวจหลักฐานรายได้และกิจการ"] + common("CHECKLIST"),
        ["AG02", "อาชีพอิสระจดทะเบียน / เจ้าของกิจการ", "ผู้กู้มีทะเบียนการค้า/ทะเบียนพาณิชย์/นิติบุคคลเกี่ยวข้อง", "ตรวจเอกสารจดทะเบียนและสิทธิในกิจการ", "ต้องตรวจสัดส่วน/บทบาทและบัญชีหมุนเวียน"] + common("CHECKLIST"),
        ["AG03", "อาชีพอิสระ Premium", "กลุ่มวิชาชีพ/ธุรกิจที่เข้าเกณฑ์ Premium", "ตรวจเงื่อนไข Premium ก่อนใช้", "ต้องมีเอกสารเฉพาะกลุ่ม Premium"] + common("CHECKLIST"),
    ],
)

# Loan purposes
purposes = [
    ["PUR001", "ซื้อที่ดินพร้อมอาคาร"],
    ["PUR002", "ซื้อห้องชุด"],
    ["PUR003", "ซื้อที่ดิน"],
    ["PUR004", "ซื้อที่ดินพร้อมปลูกสร้างอาคาร"],
    ["PUR005", "ปลูกสร้าง"],
    ["PUR006", "ต่อเติม / ซ่อมแซมอาคาร"],
    ["PUR007", "ไถ่ถอนจำนอง"],
    ["PUR008", "ชำระหนี้เกี่ยวกับที่อยู่อาศัย"],
    ["PUR009", "ซื้อหรือจัดให้มีอุปกรณ์หรือสิ่งอำนวยความสะดวก"],
    ["PUR010", "ชำระค่าเบี้ยประกันชีวิตคุ้มครองสินเชื่อ"],
]
add_sheet("22_Loan_Purpose_Master_10", ["purpose_id", "purpose_name", "aliases", "purpose_group", "description"] + COMMON, [[p[0], p[1], "", "Housing Loan Purpose", "วัตถุประสงค์การให้กู้ตามหลักเกณฑ์"] + common("CHECKLIST") for p in purposes])

# Documents / checklist placeholder sheets
for sheet_name, headers in {
    "23_Personal_Documents": ["document_id", "document_group", "document_name", "required_level", "condition", "customer_text", "staff_note"],
    "24_Purpose_Documents": ["purpose_document_id", "purpose_id", "document_group", "document_name", "required_level", "condition", "customer_text", "staff_note", "risk_if_missing"],
    "25_Checklist_Rules_28Occ": ["checklist_rule_id", "applicant_group_id", "occupation_id", "purpose_id", "include_personal_docs", "include_income_docs", "include_purpose_docs", "additional_documents", "staff_note", "customer_note"],
    "26_Risk_Interview_Focus": ["risk_focus_id", "applicant_group_id", "occupation_id", "risk_flag", "interview_question", "staff_action", "show_to_customer"],
}.items():
    add_sheet(sheet_name, headers + COMMON, [])

# PDF templates
add_sheet(
    "28_Customer_PDF_Template",
    ["template_id", "template_name", "paper_size", "orientation", "include_customer_name", "include_staff_contact", "include_risk_flags", "include_interview_focus", "include_source_reference", "include_signature"] + COMMON,
    [["PDF_CUSTOMER_CHECKLIST_V1", "Checklist เอกสารสำหรับลูกค้า", "A4", "portrait", True, True, False, False, False, True] + common("PDF")],
)
add_sheet(
    "29_Staff_Checklist_Template",
    ["template_id", "template_name", "paper_size", "orientation", "include_risk_flags", "include_interview_focus", "include_source_reference", "include_staff_internal_note"] + COMMON,
    [["PDF_STAFF_CHECKLIST_V1", "Checklist ภายในสำหรับพนักงาน", "A4", "portrait", True, True, True, True] + common("PDF")],
)
field_map = [
    ["OF001", "customer_name", "ลูกค้าผู้มาติดต่อ", True, True],
    ["OF002", "contact_info", "เบอร์/LINE ลูกค้า", False, True],
    ["OF003", "created_at", "วันที่สร้าง", True, True],
    ["OF004", "occupation_name", "อาชีพ", True, True],
    ["OF005", "purpose_name", "วัตถุประสงค์การกู้", True, True],
    ["OF006", "branch_name", "สาขา", True, True],
    ["OF007", "branch_phone", "โทรศัพท์สาขา", True, True],
    ["OF008", "risk_flags", "Risk Flag", False, True],
    ["OF009", "interview_focus", "Interview Focus", False, True],
    ["OF010", "source_reference", "Source Reference", False, True],
]
add_sheet("29A_Output_Field_Map", ["output_field_id", "source_field", "display_label", "show_in_customer_pdf", "show_in_staff_pdf"] + COMMON, [row + common("PDF") for row in field_map])
add_sheet(
    "29B_PDF_Disclaimer",
    ["disclaimer_id", "target", "disclaimer_text"] + COMMON,
    [["PDFD001", "Customer PDF", "เอกสารนี้เป็นรายการตรวจสอบเบื้องต้นเพื่อช่วยเตรียมเอกสารประกอบการขอสินเชื่อเท่านั้น การพิจารณาสินเชื่อเป็นไปตามหลักเกณฑ์และเงื่อนไขของธนาคาร และธนาคารอาจขอเอกสารเพิ่มเติมตามประเภทสินเชื่อ วัตถุประสงค์การกู้ และข้อมูลประกอบของลูกค้า เอกสารนี้ไม่ใช่ผลการอนุมัติสินเชื่อ และไม่ถือเป็นการรับรองวงเงินกู้"] + common("PDF")],
)
add_sheet("29C_PDF_QA_Test", ["test_id", "test_case", "pass_criteria"] + COMMON, [
    ["PDFT001", "Customer PDF data leakage", "ไม่มี Risk Flag / Interview Focus / Source Reference / เบอร์ลูกค้า / LINE ID"] + common("VALIDATION"),
    ["PDFT002", "Staff Checklist internal fields", "มี Risk Flag / Interview Focus / Source Reference"] + common("VALIDATION"),
])

# Learning and Q&A
learning_topics = [
    ["LG001", "วิธีเลือกกลุ่มผู้กู้ 3 ส่วน", "Quick Guide", "AG01/AG02/AG03"],
    ["LG002", "วิธีเลือก OCC001-OCC028", "Quick Guide", "ใช้ occupation master"],
    ["LG003", "วิธีเลือก PUR001-PUR010", "Quick Guide", "ใช้ purpose master"],
    ["LG004", "การใช้ผลคำนวณโดยไม่ให้เข้าใจผิดว่าอนุมัติ", "Common Mistakes", "ต้องมี Disclaimer"],
]
add_sheet("30_Learning_Content", ["learning_id", "title", "category", "key_takeaway"] + COMMON, [row + common("LEARNING") for row in learning_topics])
add_sheet("31_Quick_Guide", ["guide_id", "title", "content"] + COMMON, [["QG001", "เริ่มใช้งาน 4 เมนู", "เลือกเมนูตามงาน: คำนวณ / Checklist / Self Learning / Q&A"] + common("LEARNING")])
add_sheet("32_Common_Mistakes", ["mistake_id", "topic", "mistake_description", "correct_practice"] + COMMON, [["CM001", "AI เดาข้อมูล", "ปล่อยให้ AI สร้างเกณฑ์เอง", "ตอบเฉพาะข้อมูล Approved"] + common("LEARNING")])
add_sheet("33_Infographic_Index", ["infographic_id", "title", "topic", "source_file"] + COMMON, [["INFO001", "Self Welfare", "Policy communication", "uploaded infographic"] + common("LEARNING", review_status="Need Review")])
add_sheet("35_Learning_Quiz", ["quiz_id", "question", "option_a", "option_b", "option_c", "option_d", "correct_answer", "explanation"] + COMMON, [["QZ001", "AI ตอบได้จากแถวใด", "Draft", "Need Review", "Approved", "Inactive", "C", "อ่านเฉพาะ Approved"] + common("LEARNING")])

qa_categories = [
    ["QA01", "เอกสารรายได้"], ["QA02", "อาชีพอิสระ"], ["QA03", "วัตถุประสงค์การกู้"], ["QA04", "Statement / ภาระหนี้"], ["QA05", "Self Welfare"], ["QA06", "พนักงานเจ้าของโครงการ / Conflict of Interest"], ["QA07", "GHB ALL GEN / PDPA"], ["QA08", "Market Conduct"], ["QA09", "ข้อผิดพลาดที่พบบ่อย"], ["QA10", "Escalation / ต้องส่งต่อ"],
]
add_sheet("41_QA_Category", ["category_id", "category_name"] + COMMON, [row + common("QA") for row in qa_categories])
add_sheet("40_QA_Master", ["qa_id", "category_id", "question", "normalized_question", "answer", "short_answer", "confidence_level", "escalation_required"] + COMMON, [])
add_sheet("42_QA_Source_Map", ["source_map_id", "qa_id", "source_document", "source_page", "source_section", "source_summary", "source_priority"] + COMMON, [])
add_sheet("43_QA_Fallback", ["fallback_id", "fallback_text"] + COMMON, [["FB001", "ยังไม่พบข้อมูลในฐานความรู้ที่ผ่านการตรวจสอบ กรุณาตรวจสอบคู่มืออย่างเป็นทางการ หรือส่งต่อผู้รับผิดชอบเพื่อพิจารณาเพิ่มเติม"] + common("QA")])
add_sheet("44_QA_Escalation", ["escalation_id", "condition", "action"] + COMMON, [["ESC001", "คำถามเกินฐาน Approved", "ส่งต่อผู้รับผิดชอบ"] + common("QA")])

# QA test sheets
for sheet in ["17_Calc_QA_Test", "27_Checklist_QA_Test", "36_Learning_QA_Test", "45_QA_Test", "80_QA_Test_4Menu"]:
    add_sheet(sheet, ["test_id", "scenario", "input_summary", "expected_output", "pass_criteria"] + COMMON, [])

validation_rows = [
    ["VAL001", "มี 4 เมนูครบ", "PASS if MENU01-MENU04 exist"],
    ["VAL002", "มี AG01-AG03 ครบ", "PASS if 3 applicant groups exist"],
    ["VAL003", "มี PUR001-PUR010 ครบ", "PASS if 10 purposes exist"],
    ["VAL004", "มี OCC001-OCC028 ครบ", "PASS if 28 occupation IDs exist; names may remain Need Review until mapped"],
    ["VAL005", "ไม่มี Draft/Need Review ถูก runtime ใช้ตอบ", "PASS if runtime loader filters Approved only"],
    ["VAL006", "Customer PDF ไม่มีข้อมูลภายใน", "PASS if data leakage test passed"],
    ["VAL007", "AI ไม่เดา % ค่าใช้จ่าย", "PASS if missing rate asks for actual expense"],
]
add_sheet("81_Data_Validation_Checklist", ["validation_id", "check_item", "pass_criteria"] + COMMON, [row + common("VALIDATION") for row in validation_rows])
add_sheet("99_Change_Log", ["change_id", "change_date", "version", "changed_area", "description", "changed_by", "status"] + COMMON, [["CHG001", REVIEW_DATE, VERSION, "Initial v3 blueprint", "Create Runtime KB v3 structure", REVIEWED_BY, "Draft created"] + common("CORE")])


def create_workbook() -> Path:
    wb = Workbook()
    default = wb.active
    wb.remove(default)

    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    thin = Side(style="thin", color="D9E2F3")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for sheet_name, spec in SHEETS.items():
        ws = wb.create_sheet(title=sheet_name[:31])
        headers = spec["headers"]
        rows = spec["rows"]
        ws.append(headers)
        for row in rows:
            ws.append(row)

        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border = border
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions
        for col_idx, column_cells in enumerate(ws.columns, 1):
            max_len = max(len(str(cell.value)) if cell.value is not None else 0 for cell in column_cells)
            ws.column_dimensions[get_column_letter(col_idx)].width = min(max(max_len + 2, 12), 45)

        # Basic validations on common fields
        headers_map = {h: idx + 1 for idx, h in enumerate(headers)}
        if "review_status" in headers_map:
            col = get_column_letter(headers_map["review_status"])
            dv = DataValidation(type="list", formula1='"Draft,Need Review,Approved,Inactive"', allow_blank=False)
            ws.add_data_validation(dv)
            dv.add(f"{col}2:{col}500")
        if "phase" in headers_map:
            col = get_column_letter(headers_map["phase"])
            dv = DataValidation(type="list", formula1='"Phase1,Phase2"', allow_blank=False)
            ws.add_data_validation(dv)
            dv.add(f"{col}2:{col}500")
        if "menu_scope" in headers_map:
            col = get_column_letter(headers_map["menu_scope"])
            dv = DataValidation(type="list", formula1='"CORE,CALC,CHECKLIST,LEARNING,QA,GUARDRAIL,PDF,BRANCH,VALIDATION"', allow_blank=False)
            ws.add_data_validation(dv)
            dv.add(f"{col}2:{col}500")

    output_path = OUTPUT_DIR / "Homie_KB_Runtime_Phase1_4Menu_ManualAligned_v3.xlsx"
    wb.save(output_path)
    return output_path


def create_docx() -> Path | None:
    if Document is None:
        print("python-docx is not installed. Skipping docx generation.")
        return None

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "TH Sarabun New"
    styles["Normal"].font.size = Pt(14)

    doc.add_heading("Developer Brief — Homie Credit Assistant on LINE OA Phase 1", level=1)
    doc.add_paragraph("เอกสารนี้ใช้ส่งผู้พัฒนาระบบเพื่อสร้าง Prototype และ Runtime Knowledge Base v3")

    sections = [
        ("Executive Summary", "ระบบ Homie Credit Assistant เป็นผู้ช่วยพนักงานสินเชื่อสาขา ไม่ใช่ระบบอนุมัติสินเชื่อ"),
        ("Locked 4 Menus", "1) คำนวณสินเชื่ออาชีพอิสระ 2) Checklist 3) Self Learning 4) Q&A"),
        ("Core Governance", "Runtime KB ต้องสะอาด / AI ห้ามเดาข้อมูลสินเชื่อเอง / ข้อมูลลูกค้าต้องแยกจากฐานความรู้"),
        ("Runtime Read Rule", "ระบบอ่านเฉพาะ is_active=TRUE, phase=Phase1, review_status=Approved"),
        ("Menu 1", "คำนวณรายได้สุทธิ ภาระหนี้ เงินงวดเบื้องต้น และ Estimated DSR โดยไม่ฟันธงผลอนุมัติ"),
        ("Menu 2", "สร้าง Checklist จาก AG + OCC + PUR และแยก Customer PDF กับ Staff Checklist"),
        ("Menu 3", "Quick Guide, Common Mistakes, Infographic Index, Quiz"),
        ("Menu 4", "ตอบ Q&A จากฐาน Approved เท่านั้น หากไม่พบให้ fallback/escalate"),
        ("Customer PDF", "ห้ามแสดง Risk Flag, Interview Focus, Staff Note, Source Reference, เบอร์ลูกค้า, LINE ID, รายได้หรือ DSR"),
        ("Acceptance Criteria", "ระบบต้องผ่าน QA 4 เมนูและไม่ใช้ข้อมูล Draft/Need Review ใน runtime answer"),
    ]
    for title, body in sections:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)

    output_path = OUTPUT_DIR / "Homie_Developer_Brief_Phase1_4Menu_v3.docx"
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    xlsx = create_workbook()
    print(f"Created: {xlsx}")
    docx = create_docx()
    if docx:
        print(f"Created: {docx}")
